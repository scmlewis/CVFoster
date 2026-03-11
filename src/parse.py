"""
CV Parsing module.
Handles extraction and parsing of PDF, DOCX, and TXT files.
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

# Common CV section headers (for heuristic section detection)
SECTION_HEADERS = {
    'experience': r'\b(professional\s+experience|work\s+experience|employment|experience)\b',
    'education': r'\b(education|academic|degrees|qualifications)\b',
    'skills': r'\b(technical\s+skills|skills|competencies|technical\s+competencies)\b',
    'summary': r'\b(professional\s+summary|summary|objective|profile)\b',
    'certifications': r'\b(certifications|licenses|credentials)\b',
    'projects': r'\b(projects|portfolio|selected\s+projects)\b',
}

class CVParser:
    """Handler for parsing CV documents in multiple formats."""

    @staticmethod
    def parse_pdf(file_path: str, use_ocr: bool = False) -> Tuple[str, Dict]:
        """
        Parse PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            use_ocr: If True, use OCR for scanned PDFs (fallback)
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            page_count = len(doc)

            for page_num, page in enumerate(doc):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()

            doc.close()

            metadata = {
                'source': file_path,
                'format': 'pdf',
                'pages': page_count,
                'extraction_method': 'pdf_parser'
            }

            return text, metadata

        except Exception as e:
            logger.error(f"PDF parsing failed for {file_path}: {e}")
            if use_ocr:
                logger.info("Attempting OCR fallback...")
                return CVParser._parse_pdf_ocr(file_path)
            raise

    @staticmethod
    def _parse_pdf_ocr(file_path: str) -> Tuple[str, Dict]:
        """
        Fallback OCR extraction for scanned PDFs.
        Requires Tesseract installed.
        """
        try:
            import pytesseract
            from PIL import Image
            
            doc = fitz.open(file_path)
            text = ""
            page_count = len(doc)

            for page_num, page in enumerate(doc):
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                text += f"\n--- Page {page_num + 1} (OCR) ---\n{ocr_text}"

            doc.close()

            metadata = {
                'source': file_path,
                'format': 'pdf',
                'pages': page_count,
                'extraction_method': 'ocr'
            }

            return text, metadata

        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise

    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, Dict]:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = Document(file_path)
            text = ""

            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            metadata = {
                'source': file_path,
                'format': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'extraction_method': 'docx_parser'
            }

            return text, metadata

        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise

    @staticmethod
    def parse_txt(file_path: str, encoding: str = 'utf-8') -> Tuple[str, Dict]:
        """
        Parse TXT file and extract text.
        
        Args:
            file_path: Path to TXT file
            encoding: File encoding (default: utf-8)
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()

            metadata = {
                'source': file_path,
                'format': 'txt',
                'encoding': encoding,
                'extraction_method': 'txt_parser'
            }

            return text, metadata

        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed for {file_path}, retrying with latin-1")
            return CVParser.parse_txt(file_path, encoding='latin-1')

        except Exception as e:
            logger.error(f"TXT parsing failed for {file_path}: {e}")
            raise

    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict]:
        """
        Auto-detect file format and parse accordingly.
        
        Args:
            file_path: Path to file
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        path = Path(file_path)
        file_ext = path.suffix.lower()

        if file_ext == '.pdf':
            return CVParser.parse_pdf(file_path, use_ocr=True)
        elif file_ext == '.docx':
            return CVParser.parse_docx(file_path)
        elif file_ext in ['.txt', '.text']:
            return CVParser.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: .pdf, .docx, .txt")

    @staticmethod
    def detect_sections(text: str) -> Dict[str, Tuple[int, int]]:
        """
        Detect CV sections using heuristic pattern matching.
        
        Args:
            text: Extracted CV text
        
        Returns:
            Dict mapping section name to (start_char, end_char) offsets
        """
        sections = {}
        text_lower = text.lower()
        all_matches = []

        # Find all section header matches with their positions and section types
        for section_name, pattern in SECTION_HEADERS.items():
            for m in re.finditer(pattern, text_lower):
                all_matches.append({
                    'start': m.start(),
                    'end': m.end(),
                    'section': section_name,
                    'match_text': text[m.start():m.end()]
                })

        # Sort by position
        all_matches.sort(key=lambda x: x['start'])
        
        match_info = ', '.join([f"{m['section']}@{m['start']}" for m in all_matches])
        logger.debug(f"Found {len(all_matches)} section header matches: {match_info}")

        # Build section boundaries
        # Each section starts at its header and ends at the next header
        for i, match in enumerate(all_matches):
            section_name = match['section']
            start_pos = match['start']
            
            # Find the end position (start of next section)
            if i + 1 < len(all_matches):
                end_pos = all_matches[i + 1]['start']
            else:
                end_pos = len(text)
            
            # Only store the FIRST occurrence of each section name
            # (in case there are duplicates)
            if section_name not in sections:
                sections[section_name] = (start_pos, end_pos)
                logger.debug(f"Mapped section '{section_name}' to positions {start_pos}-{end_pos}")

        return sections

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract CV sections as text, removing section headers but preserving content.
        
        Args:
            text: Full CV text
        
        Returns:
            Dict mapping section name to section text (without header)
        """
        sections_offsets = CVParser.detect_sections(text)
        sections_text = {}

        for section_name, (start, end) in sections_offsets.items():
            # Get the exact section text between boundaries
            section_raw = text[start:end]
            
            # Find the end of the first line (the header)
            first_newline = section_raw.find('\n')
            
            if first_newline != -1:
                # Skip the header line (everything before first newline)
                content = section_raw[first_newline + 1:]
            else:
                # If no newline found, the entire section is just the header
                content = ""
            
            # Strip leading and trailing whitespace
            content = content.strip()
            
            # Only add sections that have real content
            if content and len(content) > 5:
                sections_text[section_name] = content
                logger.debug(f"Extracted section '{section_name}': {len(content)} chars")
            else:
                logger.debug(f"Skipped section '{section_name}': insufficient content")

        logger.info(f"Successfully extracted {len(sections_text)} sections: {list(sections_text.keys())}")
        return sections_text
