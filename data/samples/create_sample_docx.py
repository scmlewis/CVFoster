#!/usr/bin/env python
"""
Generate sample DOCX CV files for testing
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_docx():
    """Create a sample DOCX CV file"""
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('EMILY RODRIGUEZ')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact = doc.add_paragraph()
    contact_run = contact.add_run('emily.rodriguez@email.com | (555) 321-7654 | LinkedIn: linkedin.com/in/emilyrodriguez | Seattle, WA')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Professional Summary
    summary_heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    summary_text = doc.add_paragraph(
        'Full-Stack Web Developer with 5+ years of experience building modern, responsive web applications. '
        'Expertise in React, Node.js, and cloud deployment. Passionate about creating intuitive user interfaces '
        'and writing clean, maintainable code. Strong problem-solving skills with proven ability to lead small teams.'
    )
    
    # Professional Experience
    exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
    
    # Job 1
    job1_title = doc.add_paragraph('Senior Frontend Developer', style='List Bullet')
    job1_title.clear()
    job1_run = job1_title.add_run('Senior Frontend Developer')
    job1_run.bold = True
    job1_company = doc.add_paragraph('Web Agency Pro, Seattle, WA | April 2021 - Present', style='List Bullet')
    job1_company.clear()
    job1_company.add_run('Web Agency Pro, Seattle, WA | April 2021 - Present').italic = True
    
    doc.add_paragraph('Led frontend development for 10+ client projects using React and TypeScript', style='List Bullet')
    doc.add_paragraph('Improved application performance by 45% through code optimization and lazy loading', style='List Bullet')
    doc.add_paragraph('Mentored 3 junior developers on best practices and modern web development', style='List Bullet')
    doc.add_paragraph('Implemented responsive design frameworks serving 500K+ monthly users', style='List Bullet')
    
    # Job 2
    job2_run = doc.add_paragraph('Full-Stack Developer').runs[0]
    job2_run.bold = True
    job2_company = doc.add_paragraph('Digital Solutions LLC, Portland, OR | July 2019 - March 2021').runs[0]
    job2_company.italic = True
    
    doc.add_paragraph('Developed full-stack web applications using React, Node.js, and MongoDB', style='List Bullet')
    doc.add_paragraph('Designed and implemented REST APIs serving 100K+ daily active users', style='List Bullet')
    doc.add_paragraph('Deployed applications to AWS using Docker and Elastic Beanstalk', style='List Bullet')
    
    # Education
    edu_heading = doc.add_heading('EDUCATION', level=2)
    
    degree = doc.add_paragraph('Bachelor of Science in Computer Science')
    degree.runs[0].bold = True
    school = doc.add_paragraph('Pacific University, Portland, OR | Graduated May 2019')
    school.runs[0].italic = True
    
    # Skills
    skills_heading = doc.add_heading('TECHNICAL SKILLS', level=2)
    
    skills_text = (
        'Frontend: React, JavaScript, TypeScript, HTML5, CSS3, Webpack, Jest\n'
        'Backend: Node.js, Express, Python, Flask, RESTful APIs\n'
        'Databases: MongoDB, PostgreSQL, Firebase\n'
        'Cloud & DevOps: AWS, Docker, Git, GitHub Actions\n'
        'Tools: VS Code, Figma, JIRA, Postman'
    )
    doc.add_paragraph(skills_text)
    
    # Save document
    output_path = os.path.join(
        os.path.dirname(__file__),
        'Sample_CV_4_Emily_Rodriguez.docx'
    )
    
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_sample_docx()
    print("\nSample DOCX CV created successfully!")
